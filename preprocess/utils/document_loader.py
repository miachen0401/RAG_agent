"""
Document Loader and Section Extractor

This module handles loading documents (Word .docx and .txt files) and
extracting sections from them based on heading detection.
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl


class DocumentLoader:
    """
    Loader for Word documents and text files with section extraction.
    """

    def __init__(self):
        """Initialize the document loader."""
        pass

    def extract_eln_id(self, text: str) -> Optional[str]:
        """
        Extract ELN ID from document text if present.

        Looks for patterns like "ELN-12345" or "ELN_12345" or similar.

        Args:
            text: Document text to search

        Returns:
            ELN ID if found, None otherwise
        """
        # Common ELN ID patterns
        patterns = [
            r'ELN[-_]?\d+',
            r'eln[-_]?\d+',
            r'ELN\s*ID\s*:?\s*(\w+)',
            r'eln\s*id\s*:?\s*(\w+)'
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(0) if '(' not in pattern else match.group(1)

        return None

    def load_text_file(self, file_path: str) -> Dict[str, Any]:
        """
        Load a plain text file and extract sections.

        Sections are detected by lines that look like headers:
        - Lines ending with ":"
        - Lines that are all caps
        - Lines that start with numbers (1., 2., etc.)
        - Lines that start with # (markdown-style)

        Args:
            file_path: Path to the text file

        Returns:
            Dictionary with document metadata and sections
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # Extract ELN ID
        eln_id = self.extract_eln_id(content)

        # Split into lines
        lines = content.split('\n')

        sections = []
        current_section = {
            "section_id": 0,
            "section_name": "Introduction",
            "text": ""
        }
        section_id = 0

        for line in lines:
            stripped = line.strip()

            # Check if this line is a section header
            is_header = False
            if stripped:
                # Check various header patterns
                if (
                    # Ends with colon
                    stripped.endswith(':') or
                    # All caps and short (likely a header)
                    (stripped.isupper() and len(stripped.split()) <= 8) or
                    # Starts with markdown header
                    stripped.startswith('#') or
                    # Starts with number (1., 2.1, etc.)
                    re.match(r'^\d+\.(\d+\.)*\s+\w+', stripped)
                ):
                    is_header = True

            if is_header and current_section["text"].strip():
                # Save the current section
                sections.append(current_section)

                # Start new section
                section_id += 1
                section_name = stripped.rstrip(':').lstrip('#').strip()
                current_section = {
                    "section_id": section_id,
                    "section_name": section_name,
                    "text": ""
                }
            else:
                # Add line to current section
                current_section["text"] += line + "\n"

        # Add the last section
        if current_section["text"].strip():
            sections.append(current_section)

        return {
            "file_path": file_path,
            "eln_id": eln_id,
            "sections": sections
        }

    def load_docx_file(self, file_path: str) -> Dict[str, Any]:
        """
        Load a Word .docx file and extract sections based on heading styles.

        Args:
            file_path: Path to the .docx file

        Returns:
            Dictionary with document metadata and sections
        """
        doc = Document(file_path)

        # Extract all text for ELN ID detection
        full_text = "\n".join([para.text for para in doc.paragraphs])
        eln_id = self.extract_eln_id(full_text)

        sections = []
        current_section = {
            "section_id": 0,
            "section_name": "Introduction",
            "text": ""
        }
        section_id = 0

        for para in doc.paragraphs:
            # Check if this is a heading (based on style name)
            is_heading = False
            if para.style.name.startswith('Heading'):
                is_heading = True
            # Also check for common heading patterns in text
            elif para.text.strip():
                stripped = para.text.strip()
                if (
                    stripped.endswith(':') or
                    (stripped.isupper() and len(stripped.split()) <= 8) or
                    re.match(r'^\d+\.(\d+\.)*\s+\w+', stripped)
                ):
                    is_heading = True

            if is_heading and para.text.strip():
                # Save current section if it has content
                if current_section["text"].strip():
                    sections.append(current_section)

                # Start new section
                section_id += 1
                section_name = para.text.strip().rstrip(':')
                current_section = {
                    "section_id": section_id,
                    "section_name": section_name,
                    "text": ""
                }
            else:
                # Add paragraph to current section
                if para.text.strip():
                    current_section["text"] += para.text + "\n"

        # Add the last section
        if current_section["text"].strip():
            sections.append(current_section)

        # Handle tables
        for table_idx, table in enumerate(doc.tables):
            table_text = []
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_text))

            if table_text:
                sections.append({
                    "section_id": section_id + 1 + table_idx,
                    "section_name": f"Table {table_idx + 1}",
                    "text": "\n".join(table_text)
                })

        return {
            "file_path": file_path,
            "eln_id": eln_id,
            "sections": sections
        }

    def load_document(self, file_path: str) -> Dict[str, Any]:
        """
        Load a document (auto-detects type based on extension).

        Args:
            file_path: Path to the document file

        Returns:
            Dictionary with document metadata and sections

        Raises:
            ValueError: If file type is not supported
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = file_path.suffix.lower()

        if suffix == '.txt':
            return self.load_text_file(str(file_path))
        elif suffix in ['.docx', '.doc']:
            if suffix == '.doc':
                raise ValueError(
                    ".doc files are not supported. Please convert to .docx format."
                )
            return self.load_docx_file(str(file_path))
        else:
            raise ValueError(f"Unsupported file type: {suffix}")


def find_documents_in_folders(root_dir: str) -> List[Dict[str, str]]:
    """
    Find documents in folder structure.

    For each folder under root_dir, looks for a Word document (.docx)
    directly under that folder (not in subfolders).

    Args:
        root_dir: Root directory to search (e.g., "sample_documents")

    Returns:
        List of dictionaries with 'folder_name' and 'file_path'
    """
    root_path = Path(root_dir)
    documents = []

    # check subdirectories (each folder should have one doc directly under it)
    for folder in root_path.iterdir():
        if not folder.is_dir():
            continue

        # Look for .docx files directly under this folder
        docx_files = list(folder.glob("*.docx"))
        if docx_files:
            documents.append({
                "folder_name": folder.name,
                "file_path": str(docx_files[0])  # Take the first one
            })
            continue

    return documents
